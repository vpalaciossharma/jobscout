"""
resume_tailor.py — Job Scout Resume Tailoring Agent
----------------------------------------------------
Reads top_jobs.json produced by job_scout.py, calls the Claude API to tailor
Vanessa's resume for each top-scoring job, generates a Word doc (.docx) per
job, and emails them as attachments.

Requires: anthropic, python-docx (see requirements.txt)
Secrets:  ANTHROPIC_API_KEY, GMAIL_ADDRESS, GMAIL_APP_PASSWORD, RECIPIENT_EMAIL
"""

import json
import os
import re
import smtplib
import sys
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────

SCRIPT_DIR = Path(__file__).parent
TOP_JOBS_PATH = SCRIPT_DIR / "top_jobs.json"
RESUME_PATH = SCRIPT_DIR / "resume.txt"
CONFIG_PATH = SCRIPT_DIR / "config.json"
OUTPUT_DIR = SCRIPT_DIR / "tailored_resumes"

# Only tailor resumes for the top N jobs to keep email size manageable
MAX_RESUMES = 5

CLAUDE_MODEL = "claude-haiku-4-5-20251001"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def load_config() -> dict:
    with open(CONFIG_PATH) as f:
        return json.load(f)


def load_resume() -> str:
    with open(RESUME_PATH) as f:
        return f.read()


def load_top_jobs() -> list:
    if not TOP_JOBS_PATH.exists():
        print("[WARN] top_jobs.json not found — run job_scout.py first.")
        return []
    with open(TOP_JOBS_PATH) as f:
        return json.load(f)


def safe_filename(job_title: str, employer: str) -> str:
    """Create a safe filename from job title and employer."""
    raw = f"VPalaciosSharma_{employer}_{job_title}"
    safe = re.sub(r"[^\w\s-]", "", raw).strip()
    safe = re.sub(r"[\s]+", "_", safe)
    return safe[:80] + ".docx"


# ─────────────────────────────────────────────────────────────────────────────
# CLAUDE API — RESUME TAILORING
# ─────────────────────────────────────────────────────────────────────────────

TAILOR_PROMPT = """You are an expert resume writer helping a Senior Technical Product Manager tailor her resume for a specific job.

JOB DETAILS:
Title: {job_title}
Company: {employer}
Score: {score}/100 (how well her background matches this role)
Job Description (truncated):
{jd}

CURRENT RESUME:
{resume}

Your task: Tailor this resume to maximize relevance for this specific role. Be strategic but keep everything factually accurate.

Return ONLY a valid JSON object — no markdown, no explanation — with exactly these keys:

{{
  "tailored_summary": "Rewritten professional summary (3-4 sentences). Lead with the most relevant expertise for this role. Naturally incorporate 2-3 keywords from the JD.",
  "key_skills": "Reordered/filtered skills line most relevant to this JD (keep the same format: skill1 | skill2 | skill3 ...)",
  "experience": [
    {{
      "company": "Company name",
      "role": "Job title",
      "period": "Date range",
      "bullets": ["bullet 1", "bullet 2", ...]
    }}
  ],
  "change_summary": [
    "Summary line 1: what changed and why (max 15 words)",
    "Summary line 2: ...",
    "Summary line 3: ..."
  ]
}}

Rules for the experience section:
- Include ALL roles from the original resume
- Keep all bullets factually identical — only REORDER them to put most relevant first
- You may lightly rephrase a bullet to swap in a JD keyword IF the meaning is identical
- Keep 4-6 bullets per role (drop least relevant if needed, but never invent new ones)
- change_summary should have 3-5 items explaining the most impactful changes made
"""


def tailor_with_claude(client: anthropic.Anthropic, resume: str, job: dict) -> dict:
    """Call Claude to tailor resume for a specific job. Returns parsed dict."""
    jd = (job.get("job_description") or "")[:3500]
    prompt = TAILOR_PROMPT.format(
        job_title=job.get("job_title", ""),
        employer=job.get("employer_name", ""),
        score=job.get("score", 0),
        jd=jd,
        resume=resume,
    )

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2500,
        messages=[{"role": "user", "content": prompt}],
    )

    content = response.content[0].text.strip()

    # Strip markdown code fences if present
    if content.startswith("```"):
        content = re.sub(r"^```[a-z]*\n?", "", content)
        content = re.sub(r"\n?```$", "", content)

    return json.loads(content)


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOC GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc: Document):
    """Add a thin horizontal line to the document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def set_font(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, size=9, bold=True, color=(70, 130, 180))  # steel blue
    add_horizontal_rule(doc)


def create_word_doc(tailored: dict, job: dict) -> Document:
    """Build a clean, professional Word resume from tailored data."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Name ──────────────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run("Vanessa Palacios Sharma")
    set_font(name_run, size=18, bold=True, color=(30, 30, 30))

    # ── Contact ───────────────────────────────────────────────────────────────
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    contact_run = contact_p.add_run(
        "Portland, Oregon  ·  sharma.portland@gmail.com  ·  linkedin.com/in/vpalaciossharma/"
    )
    set_font(contact_run, size=9, color=(80, 80, 80))

    # ── Tailored for badge ────────────────────────────────────────────────────
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_p.paragraph_format.space_after = Pt(10)
    badge_run = badge_p.add_run(
        f"✦ Tailored for: {job.get('job_title','')} at {job.get('employer_name','')} ✦"
    )
    set_font(badge_run, size=8, color=(120, 120, 120))

    # ── Professional Summary ──────────────────────────────────────────────────
    section_heading(doc, "Professional Summary")
    summary_p = doc.add_paragraph(tailored.get("tailored_summary", ""))
    summary_p.paragraph_format.space_after = Pt(4)
    for run in summary_p.runs:
        set_font(run, size=10)

    # ── Skills ────────────────────────────────────────────────────────────────
    section_heading(doc, "AdTech & Key Skills")
    skills_p = doc.add_paragraph(tailored.get("key_skills", ""))
    skills_p.paragraph_format.space_after = Pt(4)
    for run in skills_p.runs:
        set_font(run, size=9.5)

    # ── Experience ────────────────────────────────────────────────────────────
    section_heading(doc, "Professional Experience")

    for exp in tailored.get("experience", []):
        # Company line
        co_p = doc.add_paragraph()
        co_p.paragraph_format.space_before = Pt(6)
        co_p.paragraph_format.space_after = Pt(1)
        co_run = co_p.add_run(exp.get("company", ""))
        set_font(co_run, size=10.5, bold=True)

        # Role + period line
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_before = Pt(0)
        role_p.paragraph_format.space_after = Pt(2)
        role_run = role_p.add_run(f"{exp.get('role','')}  |  {exp.get('period','')}")
        set_font(role_run, size=10, color=(60, 60, 60))

        # Bullets
        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.25)
            b_run = bp.add_run(bullet.lstrip("•- "))
            set_font(b_run, size=10)

    # ── Education ─────────────────────────────────────────────────────────────
    section_heading(doc, "Education")
    edu_p = doc.add_paragraph(
        "Bachelor of Arts in Psychology  |  California State University, San Marcos"
    )
    edu_p.paragraph_format.space_after = Pt(2)
    for run in edu_p.runs:
        set_font(run, size=10)

    # ── Certifications ────────────────────────────────────────────────────────
    section_heading(doc, "Certifications")
    certs = [
        "Google AI Essentials – Coursera (January 2025)",
        "PMC Level 1-3 – Pragmatic Institute (October 2021)",
        "Certified ScrumMaster (CSM) – Scrum Alliance (October 2020)",
        "Android and Web Development Certification – Epicodus (September 2016)",
    ]
    for cert in certs:
        cp = doc.add_paragraph(style="List Bullet")
        cp.paragraph_format.space_before = Pt(1)
        cp.paragraph_format.space_after = Pt(1)
        cp.paragraph_format.left_indent = Inches(0.25)
        set_font(cp.add_run(cert), size=10)

    # ── Languages ─────────────────────────────────────────────────────────────
    section_heading(doc, "Languages")
    lang_p = doc.add_paragraph("Spanish (Native)  |  English (Native)  |  Hindi (Beginner)")
    for run in lang_p.runs:
        set_font(run, size=10)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# EMAIL WITH ATTACHMENTS
# ─────────────────────────────────────────────────────────────────────────────

def build_summary_email(jobs_and_summaries: list, today_str: str) -> str:
    """Build the HTML body of the tailored resumes email."""
    job_rows = ""
    for job, summary, filename in jobs_and_summaries:
        changes_html = "".join(f"<li>{c}</li>" for c in summary)
        location = "Remote" if job.get("job_is_remote") else f"{job.get('job_city','')}, {job.get('job_state','')}"
        job_rows += f"""
        <div style="border:1px solid #e0e0e0;border-radius:8px;padding:16px;margin-bottom:16px;">
          <div style="font-size:16px;font-weight:bold;color:#1a1a2e">
            {job.get('job_title','')}
          </div>
          <div style="color:#555;font-size:13px;margin-bottom:8px;">
            {job.get('employer_name','')} · {location} · Match: {job.get('score',0)}%
          </div>
          <div style="font-size:12px;color:#333;">
            <strong>What changed:</strong>
            <ul style="margin:4px 0 0 0;padding-left:18px;">{changes_html}</ul>
          </div>
          <div style="margin-top:8px;font-size:11px;color:#888;">📎 Attached: {filename}</div>
        </div>
        """

    return f"""
    <html><body style="font-family:Arial,sans-serif;max-width:680px;margin:0 auto;padding:20px;">
      <div style="background:#1a1a2e;color:white;padding:20px;border-radius:8px;margin-bottom:24px;">
        <div style="font-size:22px;font-weight:bold;">📄 Tailored Resumes Ready</div>
        <div style="font-size:14px;opacity:.8;margin-top:4px;">{today_str} · {len(jobs_and_summaries)} resume(s) attached</div>
      </div>
      <p style="color:#444;font-size:14px;">
        Your resumes have been tailored for today's top job matches.
        Each .docx file is attached and ready to send — review before submitting!
      </p>
      {job_rows}
      <p style="font-size:11px;color:#aaa;margin-top:24px;">
        Generated by Job Scout Resume Tailor · Always review before submitting
      </p>
    </body></html>
    """


def send_email_with_attachments(config: dict, subject: str, html_body: str, attachments: list):
    """Send email with .docx file attachments."""
    msg = MIMEMultipart()
    msg["From"] = config["gmail_address"]
    msg["To"] = config["recipient_email"]
    msg["Subject"] = subject
    msg.attach(MIMEText(html_body, "html"))

    for filepath in attachments:
        with open(filepath, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename={Path(filepath).name}",
        )
        msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(config["gmail_address"], config["gmail_app_password"])
        server.send_message(msg)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    config = load_config()
    resume_text = load_resume()
    top_jobs = load_top_jobs()

    if not top_jobs:
        print("[INFO] No top jobs found — nothing to tailor. Exiting.")
        sys.exit(0)

    # Cap at MAX_RESUMES (already sorted by score from job_scout.py)
    top_jobs = top_jobs[:MAX_RESUMES]
    print(f"[Resume Tailor] Tailoring resumes for {len(top_jobs)} top job(s)...")

    # Set up output directory
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Set up Claude client
    api_key = config.get("anthropic_api_key") or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("[ERROR] No ANTHROPIC_API_KEY found in config or environment.")
        sys.exit(1)
    client = anthropic.Anthropic(api_key=api_key)

    today_str = datetime.now().strftime("%B %d, %Y")
    today_short = datetime.now().strftime("%a %b %d")

    jobs_and_summaries = []
    attachment_paths = []

    for i, job in enumerate(top_jobs, 1):
        title = job.get("job_title", "Unknown Role")
        employer = job.get("employer_name", "Unknown Company")
        print(f"  [{i}/{len(top_jobs)}] Tailoring for: {title} @ {employer} (score: {job.get('score',0)})")

        try:
            tailored = tailor_with_claude(client, resume_text, job)
        except Exception as e:
            print(f"  [WARN] Claude API call failed for {employer}: {e}")
            continue

        # Generate Word doc
        try:
            doc = create_word_doc(tailored, job)
            filename = safe_filename(title, employer)
            filepath = OUTPUT_DIR / filename
            doc.save(filepath)
            attachment_paths.append(filepath)
            jobs_and_summaries.append((job, tailored.get("change_summary", []), filename))
            print(f"  ✓ Saved: {filename}")
        except Exception as e:
            print(f"  [WARN] Word doc generation failed for {employer}: {e}")
            continue

    if not jobs_and_summaries:
        print("[WARN] No resumes were successfully generated. Skipping email.")
        sys.exit(0)

    # Send email with attachments
    print(f"\n  → Sending {len(attachment_paths)} tailored resume(s) to {config['recipient_email']}...")
    subject = f"📄 Tailored Resumes: {len(attachment_paths)} Ready · {today_short}"
    html = build_summary_email(jobs_and_summaries, today_str)

    try:
        send_email_with_attachments(config, subject, html, attachment_paths)
        print("  ✓ Email sent successfully.")
    except Exception as e:
        print(f"  [ERROR] Failed to send email: {e}")
        sys.exit(1)

    # Clean up generated files
    for p in attachment_paths:
        Path(p).unlink(missing_ok=True)
    if OUTPUT_DIR.exists():
        try:
            OUTPUT_DIR.rmdir()
        except OSError:
            pass


if __name__ == "__main__":
    main()
